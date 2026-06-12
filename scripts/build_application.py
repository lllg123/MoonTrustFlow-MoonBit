from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "competition"
TITLE = "MoonTrustFlow 项目申报书"
PROJECT = "MoonTrustFlow：MoonBit Policy-as-Code 与可信数据流治理工具包"
GITLINK = "https://www.gitlink.org.cn/lllglllg/MoonTrustFlow"
GITHUB = "https://github.com/lllg123/MoonTrustFlow-MoonBit-"
PDF = OUT / "MoonTrustFlow项目申报书.pdf"
DOCX = OUT / "MoonTrustFlow项目申报书.docx"

SECTIONS = [
    (
        "1. 项目名称",
        PROJECT,
    ),
    (
        "2. 项目简介",
        "MoonTrustFlow 是面向 MoonBit 生态的轻量策略评估工具包。项目使用 .mtf "
        "规则描述数据源、危险汇点、可信边界、净化控制点、数据流边、禁止路径、"
        "必经控制点和例外策略，并输出可解释的合规/风险报告。",
    ),
    (
        "3. 项目方向与适用场景",
        "方向为工程工具、策略评估和基础软件生态补位。适用于服务边界治理、"
        "CI 审计、架构评审、数据流控制点检查，以及后续接入 MoonBit AST 或调用图的工具链。",
    ),
    (
        "4. 拟实现的核心功能",
        "解析 .mtf 模型；支持 source/sink/sanitizer/boundary/node/edge；支持 deny、"
        "require through=、allow；支持 severity 风险等级；按图路径评估策略并生成稳定报告。",
    ),
    (
        "5. 原创性说明",
        "本项目为原创项目，不移植已有开源项目。选题避开通用数据流分析框架、"
        "依赖风险扫描和 CI 工作流检查，聚焦策略规则、可信流模型与合规报告的组合。",
    ),
    (
        "6. 仓库链接",
        f"GitHub：{GITHUB}<br/>GitLink：{GITLINK}",
    ),
    (
        "7. 提交说明",
        "仓库保留 10-20 次有效提交，提交内容覆盖项目骨架、领域模型、解析器、"
        "策略评估、CLI、测试、CI 和申报文档，不使用空提交或重复提交凑数。",
    ),
]


def font_name() -> str:
    for candidate in [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont("CN", str(candidate)))
            return "CN"
    return "Helvetica"


def build_pdf() -> None:
    font = font_name()
    styles = getSampleStyleSheet()
    title = ParagraphStyle(
        "title",
        parent=styles["Title"],
        fontName=font,
        fontSize=17,
        leading=22,
        alignment=TA_CENTER,
        textColor=colors.white,
        spaceAfter=0,
    )
    subtitle = ParagraphStyle(
        "subtitle",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=8.6,
        leading=11,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#d7efe9"),
    )
    heading = ParagraphStyle(
        "heading",
        parent=styles["Heading2"],
        fontName=font,
        fontSize=9.2,
        leading=12,
        textColor=colors.HexColor("#0f3f3a"),
        spaceBefore=0,
        spaceAfter=1,
    )
    body = ParagraphStyle(
        "body",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=8.45,
        leading=11.4,
        spaceAfter=0,
    )
    meta = ParagraphStyle(
        "meta",
        parent=styles["BodyText"],
        fontName=font,
        fontSize=8.2,
        leading=10.7,
        textColor=colors.HexColor("#132321"),
    )
    doc = SimpleDocTemplate(
        str(PDF),
        pagesize=A4,
        rightMargin=1.25 * cm,
        leftMargin=1.25 * cm,
        topMargin=1.05 * cm,
        bottomMargin=1.0 * cm,
        title=TITLE,
    )
    header = Table(
        [
            [
                Paragraph(TITLE, title),
                Paragraph(
                    "MoonBit Policy-as-Code / trusted flow governance / reproducible CI report",
                    subtitle,
                ),
            ]
        ],
        colWidths=[8.4 * cm, 9.1 * cm],
    )
    header.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#103d3a")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story = [header, Spacer(1, 0.16 * cm)]
    facts = Table(
        [
            [
                Paragraph("<b>项目</b><br/>MoonTrustFlow", meta),
                Paragraph("<b>方向</b><br/>策略评估 / 可信数据流治理", meta),
                Paragraph("<b>许可证</b><br/>Apache-2.0", meta),
                Paragraph("<b>仓库</b><br/>GitHub + GitLink 同步", meta),
            ],
        ],
        colWidths=[3.15 * cm, 6.05 * cm, 2.05 * cm, 6.25 * cm],
    )
    facts.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font),
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#e9f4f0")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#9db8b0")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.extend([facts, Spacer(1, 0.12 * cm)])
    for section_title, text in SECTIONS:
        block = Table(
            [[Paragraph(section_title, heading), Paragraph(text, body)]],
            colWidths=[3.25 * cm, 14.25 * cm],
        )
        block.setStyle(
            TableStyle(
                [
                    ("LINEBEFORE", (0, 0), (0, 0), 2.4, colors.HexColor("#2c7a66")),
                    ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#f0f7f4")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(block)
        story.append(Spacer(1, 0.055 * cm))
    doc.build(story)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(42)
    section.bottom_margin = Pt(42)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(16, 61, 58)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("MoonBit Policy-as-Code / 可信数据流治理 / 工程工具")
    subtitle_run.font.size = Pt(9)
    subtitle_run.font.color.rgb = RGBColor(44, 122, 102)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("项目", PROJECT),
        ("方向", "Policy-as-Code / 可信数据流治理 / 工程工具"),
        ("许可证", "Apache-2.0"),
        ("仓库", f"GitHub：{GITHUB}\nGitLink：{GITLINK}"),
    ]
    for row, (key, value) in zip(table.rows, rows):
        row.cells[0].text = key
        row.cells[1].text = value

    for section_title, text in SECTIONS:
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section_title)
        heading_run.bold = True
        heading_run.font.size = Pt(11)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(text.replace("<br/>", "\n"))
    doc.save(DOCX)


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    build_pdf()
    build_docx()
    print(PDF)
    print(DOCX)


if __name__ == "__main__":
    main()
